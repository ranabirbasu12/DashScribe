# exporter.py
"""Render the unified transcript payload into multiple output formats."""
import json


def _fmt_ts(seconds: float, *, srt: bool = False) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    ms = int((seconds - int(seconds)) * 1000)
    sep = "," if srt else "."
    return f"{h:02d}:{m:02d}:{s:02d}{sep}{ms:03d}"


def _fmt_clock(seconds: float) -> str:
    h = int(seconds // 3600)
    m = int((seconds % 3600) // 60)
    s = int(seconds % 60)
    return f"{h:02d}:{m:02d}:{s:02d}"


def _speaker_label(payload: dict, speaker_id: str) -> str:
    for sp in payload.get("speakers", []):
        if sp["id"] == speaker_id:
            return sp["label"]
    return speaker_id


def to_txt(payload: dict) -> str:
    return " ".join(s["text"] for s in payload["segments"]).strip()


def to_markdown(payload: dict) -> str:
    """Speaker-grouped paragraphs: consecutive same-speaker segments combine."""
    lines: list[str] = []
    current_spk: str | None = None
    current_chunks: list[str] = []
    current_start: float = 0.0
    for seg in payload["segments"]:
        if seg["speaker_id"] != current_spk:
            if current_chunks:
                lines.append(
                    f"**{_speaker_label(payload, current_spk)}** _"
                    f"{_fmt_clock(current_start)}_\n\n" + " ".join(current_chunks) + "\n"
                )
            current_spk = seg["speaker_id"]
            current_start = seg["start"]
            current_chunks = [seg["text"]]
        else:
            current_chunks.append(seg["text"])
    if current_chunks:
        lines.append(
            f"**{_speaker_label(payload, current_spk)}** _"
            f"{_fmt_clock(current_start)}_\n\n" + " ".join(current_chunks) + "\n"
        )
    return "\n".join(lines)


def to_srt(payload: dict) -> str:
    blocks: list[str] = []
    for i, seg in enumerate(payload["segments"], start=1):
        label = _speaker_label(payload, seg["speaker_id"])
        blocks.append(
            f"{i}\n"
            f"{_fmt_ts(seg['start'], srt=True)} --> {_fmt_ts(seg['end'], srt=True)}\n"
            f"{label}: {seg['text']}\n"
        )
    return "\n".join(blocks)


def to_vtt(payload: dict) -> str:
    blocks: list[str] = ["WEBVTT\n"]
    for seg in payload["segments"]:
        label = _speaker_label(payload, seg["speaker_id"])
        blocks.append(
            f"{_fmt_ts(seg['start'])} --> {_fmt_ts(seg['end'])}\n"
            f"{label}: {seg['text']}\n"
        )
    return "\n".join(blocks)


def to_json(payload: dict) -> str:
    return json.dumps(payload, indent=2, ensure_ascii=False)


def to_docx(payload: dict, dest_path: str) -> None:
    """Writes a .docx with speaker-grouped paragraphs and timestamps."""
    from docx import Document
    doc = Document()
    current_spk: str | None = None
    current_para = None
    for seg in payload["segments"]:
        if seg["speaker_id"] != current_spk:
            label = _speaker_label(payload, seg["speaker_id"])
            ts = _fmt_clock(seg["start"])
            heading = doc.add_paragraph()
            run = heading.add_run(f"{label}  ")
            run.bold = True
            heading.add_run(ts).italic = True
            current_para = doc.add_paragraph(seg["text"])
            current_spk = seg["speaker_id"]
        else:
            current_para.add_run(" " + seg["text"])
    doc.save(dest_path)


FORMATS = {
    "txt": "text/plain",
    "md": "text/markdown",
    "srt": "application/x-subrip",
    "vtt": "text/vtt",
    "json": "application/json",
    "docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}


def write_export(payload: dict, fmt: str, dest_path: str) -> None:
    if fmt not in FORMATS:
        raise ValueError(f"Unknown format: {fmt}")
    if fmt == "docx":
        to_docx(payload, dest_path)
        return
    renderer = {
        "txt": to_txt, "md": to_markdown, "srt": to_srt,
        "vtt": to_vtt, "json": to_json,
    }[fmt]
    content = renderer(payload)
    with open(dest_path, "w", encoding="utf-8") as f:
        f.write(content)
